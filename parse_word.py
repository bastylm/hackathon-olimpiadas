from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


def parse_question_bank(table):
    questions = []
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    i = 0
    while i < len(rows):
        first = rows[i][0].strip() if rows[i] else ""
        if re.match(r"^\d+\.", first):
            question = first
            i += 1
            if i < len(rows) and rows[i][0].strip().lower() == "respuesta":
                i += 1
            answers = []
            while i < len(rows):
                current = rows[i][0].strip() if rows[i] else ""
                if re.match(r"^\d+\.", current):
                    break
                if current:
                    try:
                        points = int(float(rows[i][1]))
                    except Exception:
                        points = 0
                    answers.append(
                        {
                            "text": current,
                            "points": points,
                            "justification": rows[i][2].strip() if len(rows[i]) > 2 else "",
                        }
                    )
                i += 1
            if answers:
                questions.append({"text": question, "answers": answers})
        else:
            i += 1
    return questions


def main():
    doc = Document(Path(sys.argv[1]))
    titles = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    banks = []
    for idx, table in enumerate(doc.tables):
        questions = parse_question_bank(table)
        if not questions:
            continue
        name = titles[idx] if idx < len(titles) else f"Cuestionario Word {idx + 1}"
        banks.append({"id": "", "name": name, "area": "Importado desde Word", "questions": questions})
    print(json.dumps(banks, ensure_ascii=False))


if __name__ == "__main__":
    main()
