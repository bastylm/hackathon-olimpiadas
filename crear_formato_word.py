from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


def style_run(run, bold=False, color=None, size=11):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Formato para cargar preguntas")
style_run(run, bold=True, color=(91, 33, 182), size=20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Olimpiadas Tecnologicas 2026 - Etapa Hackathon")
style_run(run, color=(114, 95, 134), size=11)

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run("Instrucciones: ").bold = True
intro.add_run(
    "complete este documento respetando la estructura. Cada pregunta debe comenzar con numero y punto. "
    "Cada alternativa debe comenzar con letra y parentesis, seguida del texto, una barra vertical y el puntaje."
)

rules = [
    "Use puntajes numericos, por ejemplo: 0, 25, 50, 75 o 100.",
    "No cambie las letras A), B), C), D) al inicio de las alternativas.",
    "Deje una linea en blanco entre preguntas.",
    "Puede agregar mas preguntas copiando el mismo formato.",
]
for item in rules:
    p = doc.add_paragraph(style=None)
    p.style = styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.2)
    p.add_run("- " + item)

doc.add_paragraph()

bank = doc.add_paragraph()
run = bank.add_run("Banco: Nombre del cuestionario o modulo")
style_run(run, bold=True, color=(91, 33, 182), size=12)

examples = [
    (
        "1. Que accion representa mejor una buena practica de seguridad digital?",
        [
            ("A) Compartir la clave con el equipo de trabajo", "0"),
            ("B) Usar la misma clave para todas las plataformas", "25"),
            ("C) Activar doble factor de autenticacion y usar claves seguras", "100"),
            ("D) Guardar la clave escrita junto al computador", "0"),
        ],
    ),
    (
        "2. Que se debe hacer antes de iniciar una actividad practica en laboratorio?",
        [
            ("A) Revisar instrucciones, riesgos y materiales necesarios", "100"),
            ("B) Comenzar rapidamente para ahorrar tiempo", "25"),
            ("C) Esperar el resultado de otros equipos", "0"),
            ("D) Apagar todos los equipos sin revisar", "0"),
        ],
    ),
]

for question, answers in examples:
    p = doc.add_paragraph()
    run = p.add_run(question)
    style_run(run, bold=True, size=11)
    for text, points in answers:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{text} | {points}")
        style_run(run, size=10.5)
    doc.add_paragraph()

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("Registro de intervencion - formato manual opcional")
style_run(run, bold=True, color=(91, 33, 182), size=16)

meta = [
    "Fecha y hora:",
    "Seccion:",
    "Banco de preguntas:",
    "Codigo:",
    "Cantidad esperada:",
]
for label in meta:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(" ______________________________________________")

doc.add_paragraph()
run = doc.add_paragraph().add_run("Cumplimiento de intervenciones")
style_run(run, bold=True, color=(91, 33, 182), size=12)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["N", "Estado", "Intervencion realizada"]
for cell, header in zip(table.rows[0].cells, headers):
    cell.text = header
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for idx in range(1, 7):
    row = table.add_row().cells
    row[0].text = str(idx)
    row[1].text = "Cumplida / Pendiente"
    row[2].text = ""

doc.add_paragraph()
run = doc.add_paragraph().add_run("Participantes y ranking")
style_run(run, bold=True, color=(91, 33, 182), size=12)
rank = doc.add_table(rows=1, cols=5)
rank.alignment = WD_TABLE_ALIGNMENT.CENTER
rank.style = "Table Grid"
for cell, header in zip(rank.rows[0].cells, ["Lugar", "Nombre", "RUT", "Respuestas", "Puntaje"]):
    cell.text = header
for _ in range(8):
    rank.add_row()

doc.save("Formato_carga_preguntas_y_registro_olimpiadas.docx")
