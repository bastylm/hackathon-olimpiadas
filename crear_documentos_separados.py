from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PURPLE = RGBColor(91, 33, 182)
MUTED = RGBColor(114, 95, 134)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = PURPLE
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(subtitle)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.color.rgb = MUTED
    doc.add_paragraph()


def set_header_cells(row, labels):
    for cell, label in zip(row.cells, labels):
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.color.rgb = PURPLE


def question_block(table, number, question, answers):
    row = table.add_row().cells
    row[0].text = f"{number}. {question}"
    row[1].text = f"{number}. {question}"
    row[2].text = f"{number}. {question}"
    set_header_cells(table.add_row(), ["Respuesta", "Pts", "Justificacion"])
    for answer, points, justification in answers:
        cells = table.add_row().cells
        cells[0].text = answer
        cells[1].text = str(points)
        cells[2].text = justification


def create_question_template():
    doc = Document()
    setup_doc(doc)
    title(doc, "Formato de carga de preguntas", "Olimpiadas Tecnologicas 2026 - Etapa Hackathon")
    doc.add_paragraph("Banco: Nombre del cuestionario o modulo")
    doc.add_paragraph(
        "Complete la tabla respetando las columnas Respuesta, Pts y Justificacion. "
        "Para agregar mas preguntas, copie el bloque completo: fila de pregunta, encabezado y alternativas."
    )
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    question_block(
        table,
        1,
        "Que caracteristica es mas importante al definir la arquitectura de conectividad del sistema?",
        [
            ("Comunicacion confiable entre sensores y controladores", 100, "La confiabilidad es critica para monitoreo y control del sistema autonomo."),
            ("Alta complejidad tecnologica", 75, "Puede aportar, pero no es el criterio principal."),
            ("Uso exclusivo de conexiones cableadas", 50, "No siempre es viable ni necesario."),
            ("Cantidad de dispositivos conectados", 25, "Es relevante, pero depende de la arquitectura."),
        ],
    )
    question_block(
        table,
        2,
        "Que protocolo seria mas adecuado para comunicacion industrial y monitoreo?",
        [
            ("Modbus TCP/IP por integracion y simplicidad", 100, "Es ampliamente utilizado en automatizacion industrial."),
            ("Bluetooth domestico", 50, "No es lo mas robusto para monitoreo industrial."),
            ("Comunicacion manual via operador", 25, "No automatiza el monitoreo."),
            ("Correo electronico automatizado", 0, "No corresponde a un protocolo de control industrial."),
        ],
    )
    doc.save("Formato_carga_preguntas.docx")


def create_report_template():
    doc = Document()
    setup_doc(doc)
    title(doc, "Reporte diario de intervencion", "Olimpiadas Tecnologicas 2026 - Etapa Hackathon")
    for label in ["Fecha", "Hora", "Seccion", "Banco de preguntas", "Codigo", "Evaluador"]:
        p = doc.add_paragraph()
        p.add_run(label + ": ").bold = True
        p.add_run("______________________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Cumplimiento de intervenciones")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_header_cells(table.rows[0], ["N", "Intervencion", "Estado", "Observacion"])
    for i in range(1, 8):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = ""
        cells[2].text = "Cumplida / Pendiente"
        cells[3].text = ""
    doc.add_paragraph()
    doc.add_paragraph("Resumen de participacion")
    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    set_header_cells(summary.rows[0], ["Esperados", "Llegaron", "Respondieron", "Pendientes"])
    summary.add_row()
    doc.add_paragraph()
    doc.add_paragraph("Ranking")
    ranking = doc.add_table(rows=1, cols=5)
    ranking.style = "Table Grid"
    set_header_cells(ranking.rows[0], ["Lugar", "Nombre", "RUT", "Respuestas", "Puntaje"])
    for _ in range(10):
        ranking.add_row()
    doc.save("Formato_reporte_intervencion_diaria.docx")


create_question_template()
create_report_template()
