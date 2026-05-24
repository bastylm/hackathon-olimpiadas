from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent
XLSX = Path(r"C:\Users\Miguel A.C.O\OneDrive\Desktop\Olimpiadas calendario.xlsx")
DOCX = Path(r"C:\Users\Miguel A.C.O\OneDrive\Desktop\Preguntas lunes 25.docx")
OUT = ROOT / "data.json"


def clean(value):
    if value is None:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def slug(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-") or "item"


def load_sections():
    wb = load_workbook(XLSX, data_only=True)
    ws = wb["Áreas"]
    headers = [clean(cell.value) for cell in ws[1]]
    sections = []
    last = {"Sectorial": "", "Área": "", "Carreras": ""}

    for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        record = dict(zip(headers, [clean(v) for v in row]))
        for key in last:
            if record.get(key):
                last[key] = record[key]
            else:
                record[key] = last[key]
        if not record.get("Sección"):
            continue
        section_id = f"{slug(record.get('Código', 'curso'))}-{slug(record['Sección'])}-{idx}"
        sections.append(
            {
                "id": section_id,
                "sectorial": record.get("Sectorial", ""),
                "area": record.get("Área", ""),
                "career": record.get("Carreras", ""),
                "code": record.get("Código", ""),
                "subject": record.get("Asignatura", ""),
                "section": record.get("Sección", ""),
                "teacher": record.get("Docente", ""),
                "date": record.get("Fecha HT", ""),
            }
        )
    return sections


def paragraph_titles(doc):
    titles = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            titles.append(text)
    return titles


def parse_question_bank(table):
    questions = []
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    i = 0
    while i < len(rows):
        first = rows[i][0].strip() if rows[i] else ""
        if re.match(r"^\d+\.", first):
            question = first
            i += 1
            if i < len(rows) and rows[i][0].lower() == "respuesta":
                i += 1
            answers = []
            while i < len(rows):
                current = rows[i][0].strip() if rows[i] else ""
                if re.match(r"^\d+\.", current):
                    break
                if current:
                    try:
                        points = int(float(rows[i][1]))
                    except Exception:
                        points = 0
                    answers.append(
                        {
                            "text": current,
                            "points": points,
                            "justification": rows[i][2].strip() if len(rows[i]) > 2 else "",
                        }
                    )
                i += 1
            questions.append({"text": question, "answers": answers})
        else:
            i += 1
    return questions


def load_banks():
    doc = Document(DOCX)
    titles = paragraph_titles(doc)
    bank_names = [
        "Telecomunicaciones y Conectividad (IEO504)",
        "Normativa Aplicada a Energías Renovables (EES401)",
        "Laboratorio de Modelado Digital (DVA302)",
    ]
    banks = []
    for idx, table in enumerate(doc.tables):
        name = bank_names[idx] if idx < len(bank_names) else f"Banco {idx + 1}"
        area = ""
        if name in titles:
            pos = titles.index(name)
            area = titles[pos - 1] if pos > 0 else ""
        banks.append(
            {
                "id": f"bank-{idx + 1}",
                "name": name,
                "area": area,
                "questions": parse_question_bank(table),
            }
        )
    return banks


def main():
    payload = {
        "source": {
            "sections": str(XLSX),
            "questions": str(DOCX),
        },
        "sections": load_sections(),
        "banks": load_banks(),
    }
    OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {OUT} with {len(payload['sections'])} sections and {len(payload['banks'])} banks")


if __name__ == "__main__":
    main()
